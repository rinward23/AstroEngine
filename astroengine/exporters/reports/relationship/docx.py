"""DOCX rendering pipeline for relationship reports."""

from __future__ import annotations

import io
import logging
import shutil
import tempfile
from collections.abc import Iterable
from pathlib import Path

from astroengine.exporters.reports.base import ReportMeta

LOG = logging.getLogger(__name__)


class DocxUnavailable(RuntimeError):
    """Raised when neither Pandoc nor the python-docx fallback can be used."""


def has_pandoc() -> bool:
    return shutil.which("pandoc") is not None


def convert_markdown_to_docx(markdown: str, meta: ReportMeta) -> bytes:
    """Convert Markdown to DOCX using Pandoc when available, otherwise python-docx."""

    try:
        if has_pandoc():
            return _convert_with_pandoc(markdown, meta)
    except Exception as exc:  # pragma: no cover - integration guard
        LOG.warning("Pandoc conversion failed, falling back to python-docx: %s", exc)
    return _convert_with_python_docx(markdown, meta)


def _convert_with_pandoc(markdown: str, meta: ReportMeta) -> bytes:
    import pypandoc

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        output_path = Path(handle.name)
    try:
        pypandoc.convert_text(
            markdown,
            to="docx",
            format="md",
            outputfile=str(output_path),
            extra_args=["--toc"],
        )
        data = output_path.read_bytes()
    finally:
        try:
            output_path.unlink(missing_ok=True)
        except OSError as exc:  # pragma: no cover - cleanup fallback
            LOG.debug("Unable to remove temporary DOCX %s: %s", output_path, exc)
    return data


def _convert_with_python_docx(markdown: str, meta: ReportMeta) -> bytes:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - library missing
        raise DocxUnavailable("python-docx not installed") from exc

    document = Document()
    _apply_meta(document, meta)
    _append_markdown_blocks(document, markdown)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _apply_meta(document: Document, meta: ReportMeta) -> None:
    core = document.core_properties
    core.title = meta.title
    core.subject = meta.subject or ""
    core.author = ", ".join(meta.authors) if meta.authors else "AstroEngine"
    core.keywords = ", ".join(meta.keywords) if meta.keywords else ""


def _append_markdown_blocks(document: Document, markdown: str) -> None:
    blocks = _split_blocks(markdown)
    for block in blocks:
        if not block:
            continue
        if block.startswith("#"):
            level = len(block) - len(block.lstrip("#"))
            text = block[level:].strip()
            document.add_heading(text, level=min(level, 4))
        elif block.startswith(('-', '*')):
            for line in block.splitlines():
                text = line.lstrip("-* ")
                document.add_paragraph(text, style="List Bullet")
        else:
            document.add_paragraph(block)


def _split_blocks(markdown: str) -> Iterable[str]:
    block: list[str] = []
    for line in markdown.splitlines():
        if not line.strip():
            if block:
                yield "\n".join(block).strip()
                block.clear()
            continue
        block.append(line.rstrip())
    if block:
        yield "\n".join(block).strip()
